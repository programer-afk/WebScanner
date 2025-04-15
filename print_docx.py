import sqlite3
import os
import argparse
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def connect_database(db_path):
    """连接到SQLite数据库"""
    try:
        conn = sqlite3.connect(db_path)
        return conn
    except sqlite3.Error as e:
        raise Exception(f"数据库连接错误: {e}")

def query_data(conn, query):
    """从数据库查询数据"""
    try:
        df = pd.read_sql_query(query, conn)
        return df
    except Exception as e:
        raise Exception(f"查询执行错误: {e}")

def format_table_as_three_line(table):
    """设置表格为三线表格式，使用XML直接设置"""
    try:
        # 使用python-docx标准API设置表格样式
        # 设置表格基本样式
        table.style = 'Table Grid'
        
        # 应用自定义格式，使用XML处理尝试实现三线表
        try:
            from docx.oxml.table import CT_Tbl
            table_element = table._element
            
            # 检查并创建tblPr元素
            tblPr = table_element.xpath('./w:tblPr')
            if not tblPr:
                tblPr = parse_xml('<w:tblPr %s></w:tblPr>' % nsdecls('w'))
                table_element.insert(0, tblPr)
            else:
                tblPr = tblPr[0]
            
            # 删除现有边框
            borders = tblPr.xpath('./w:tblBorders')
            for border in borders:
                tblPr.remove(border)
            
            # 添加新边框定义
            border_xml = f'<w:tblBorders {nsdecls("w")}>\
                <w:top w:val="single" w:sz="18" w:space="0" w:color="auto"/>\
                <w:bottom w:val="single" w:sz="18" w:space="0" w:color="auto"/>\
                <w:insideH w:val="nil"/>\
                <w:insideV w:val="nil"/>\
                <w:left w:val="nil"/>\
                <w:right w:val="nil"/>\
                </w:tblBorders>'
            borders_element = parse_xml(border_xml)
            tblPr.append(borders_element)
            
            # 设置表头单元格底部边框
            if len(table.rows) > 0:
                header_cells = table.rows[0].cells
                for cell in header_cells:
                    cell_element = cell._tc
                    tcPr = cell_element.xpath('./w:tcPr')
                    if not tcPr:
                        tcPr = parse_xml('<w:tcPr %s></w:tcPr>' % nsdecls('w'))
                        cell_element.insert(0, tcPr)
                    else:
                        tcPr = tcPr[0]
                    
                    # 删除现有边框
                    borders = tcPr.xpath('./w:tcBorders')
                    for border in borders:
                        tcPr.remove(border)
                    
                    # 添加底部边框
                    border_xml = f'<w:tcBorders {nsdecls("w")}>\
                        <w:bottom w:val="single" w:sz="8" w:space="0" w:color="auto"/>\
                        </w:tcBorders>'
                    borders_element = parse_xml(border_xml)
                    tcPr.append(borders_element)
        except Exception as e:
            print(f"警告: 无法应用完整的三线表格式: {e}")
            print("已应用基本表格样式。")
    except Exception as e:
        print(f"警告: 无法设置表格格式: {e}")
        print("继续导出基本表格...")

def create_docx_from_data(df, output_file, title="漏洞扫描结果"):
    """根据数据创建Word文档"""
    try:
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(title, level=1)

        # 创建表格
        if df.empty:
            doc.add_paragraph("没有数据")
        else:
            # 重命名列以便更好地显示
            column_map = {
                'id': '序号',
                'url': 'URL',
                'status': '状态',
                'result': '结果',
                'CVE_id': 'CVE编号',
                'time': '扫描时间'
            }
            
            # 如果列在DataFrame中存在，则重命名
            df_display = df.copy()
            for old_name, new_name in column_map.items():
                if old_name in df_display.columns:
                    df_display.rename(columns={old_name: new_name}, inplace=True)
            
            table = doc.add_table(rows=1, cols=len(df_display.columns))
            table.style = "Table Grid"

            # 添加表头
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df_display.columns):
                hdr_cells[i].text = str(col)

            # 添加数据行
            for _, row in df_display.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 应用三线表格式
            format_table_as_three_line(table)

        # 保存文档
        doc.save(output_file)
        return True
    except Exception as e:
        raise Exception(f"创建Word文档错误: {e}")

def main():
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='从SQLite数据库导出数据到Word文档三线表')
    parser.add_argument('--db', default=os.path.join(os.path.dirname(__file__), 'db.sqlite3'),
                        help='数据库文件路径')
    parser.add_argument('--query', default='SELECT * FROM vulnscan_middleware_vuln',
                        help='SQL查询语句')
    parser.add_argument('--output', default='漏洞扫描结果.docx',
                        help='输出Word文件路径')
    parser.add_argument('--title', default='Web漏洞扫描结果',
                        help='表格标题')
    args = parser.parse_args()

    try:
        # 连接数据库
        with connect_database(args.db) as conn:
            # 查询数据
            df = query_data(conn, args.query)
            
            # 创建文档
            create_docx_from_data(df, args.output, args.title)
            
            print(f"导出成功：{args.output}")
    except Exception as e:
        print(f"错误: {e}")
        return 1
    return 0

if __name__ == "__main__":
    main()