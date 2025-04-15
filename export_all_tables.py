import sqlite3
import os
import argparse
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

def connect_database(db_path):
    """连接到SQLite数据库"""
    try:
        conn = sqlite3.connect(db_path)
        return conn
    except sqlite3.Error as e:
        raise Exception(f"数据库连接错误: {e}")

def get_all_tables(conn):
    """获取数据库中所有表的列表"""
    try:
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        return [table[0] for table in tables]
    except Exception as e:
        raise Exception(f"获取表列表错误: {e}")

def query_table_data(conn, table_name):
    """查询指定表的所有数据"""
    try:
        # 检查表是否存在
        cursor = conn.cursor()
        cursor.execute(f"SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='{table_name}';")
        if cursor.fetchone()[0] == 0:
            print(f"警告: 表 '{table_name}' 不存在，跳过")
            return None
        
        # 获取表数据
        query = f"SELECT * FROM {table_name}"
        df = pd.read_sql_query(query, conn)
        return df
    except Exception as e:
        print(f"警告: 无法查询表 '{table_name}': {e}")
        return None

def format_table_as_three_line(table):
    """设置表格为三线表格式，使用XML直接设置"""
    try:
        # 使用python-docx标准API设置表格样式
        # 设置表格基本样式
        table.style = 'Table Grid'
        
        # 应用自定义格式，使用XML处理尝试实现三线表
        try:
            from docx.oxml.table import CT_Tbl
            table_element = table._element
            
            # 检查并创建tblPr元素
            tblPr = table_element.xpath('./w:tblPr')
            if not tblPr:
                tblPr = parse_xml('<w:tblPr %s></w:tblPr>' % nsdecls('w'))
                table_element.insert(0, tblPr)
            else:
                tblPr = tblPr[0]
            
            # 删除现有边框
            borders = tblPr.xpath('./w:tblBorders')
            for border in borders:
                tblPr.remove(border)
            
            # 添加新边框定义
            border_xml = f'<w:tblBorders {nsdecls("w")}>\
                <w:top w:val="single" w:sz="18" w:space="0" w:color="auto"/>\
                <w:bottom w:val="single" w:sz="18" w:space="0" w:color="auto"/>\
                <w:insideH w:val="nil"/>\
                <w:insideV w:val="nil"/>\
                <w:left w:val="nil"/>\
                <w:right w:val="nil"/>\
                </w:tblBorders>'
            borders_element = parse_xml(border_xml)
            tblPr.append(borders_element)
            
            # 设置表头单元格底部边框
            if len(table.rows) > 0:
                header_cells = table.rows[0].cells
                for cell in header_cells:
                    cell_element = cell._tc
                    tcPr = cell_element.xpath('./w:tcPr')
                    if not tcPr:
                        tcPr = parse_xml('<w:tcPr %s></w:tcPr>' % nsdecls('w'))
                        cell_element.insert(0, tcPr)
                    else:
                        tcPr = tcPr[0]
                    
                    # 删除现有边框
                    borders = tcPr.xpath('./w:tcBorders')
                    for border in borders:
                        tcPr.remove(border)
                    
                    # 添加底部边框
                    border_xml = f'<w:tcBorders {nsdecls("w")}>\
                        <w:bottom w:val="single" w:sz="8" w:space="0" w:color="auto"/>\
                        </w:tcBorders>'
                    borders_element = parse_xml(border_xml)
                    tcPr.append(borders_element)
        except Exception as e:
            print(f"警告: 无法应用完整的三线表格式: {e}")
            print("已应用基本表格样式。")
    except Exception as e:
        print(f"警告: 无法设置表格格式: {e}")
        print("继续导出基本表格...")

def add_table_to_document(doc, df, table_name):
    """将DataFrame数据添加到文档中作为三线表"""
    # 添加表格标题
    heading = doc.add_heading(level=2)
    run = heading.add_run(f"表: {table_name}")
    run.bold = True
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 检查数据
    if df is None or df.empty:
        doc.add_paragraph(f"表 '{table_name}' 没有数据")
        return False
    
    # 创建表格
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    
    # 添加数据行
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # 应用三线表格式
    format_table_as_three_line(table)
    
    # 添加空行
    doc.add_paragraph()
    
    return True

def export_all_tables_to_docx(db_path, output_file, exclude_tables=None):
    """将所有表导出到Word文档"""
    if exclude_tables is None:
        exclude_tables = ['sqlite_sequence']  # 默认排除sqlite内部表
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading("数据库表导出", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 调整页面边距为适中值
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    try:
        # 连接数据库
        with connect_database(db_path) as conn:
            # 获取所有表
            all_tables = get_all_tables(conn)
            
            # 过滤排除的表
            tables = [t for t in all_tables if t not in exclude_tables]
            
            # 添加文档摘要
            doc.add_paragraph(f"数据库包含 {len(all_tables)} 个表，本文档导出 {len(tables)} 个表")
            
            # 添加目录标题
            doc.add_heading("目录", level=2)
            
            # 添加表列表
            for i, table_name in enumerate(tables):
                para = doc.add_paragraph()
                para.add_run(f"{i+1}. {table_name}").bold = True
                
                # 创建超链接效果
                para.paragraph_format.left_indent = Cm(0.5)
            
            # 添加页面分隔符
            doc.add_page_break()
            
            # 处理每个表
            tables_processed = 0
            for table_name in tables:
                print(f"正在处理表: {table_name}")
                # 查询表数据
                df = query_table_data(conn, table_name)
                
                # 添加到文档
                if add_table_to_document(doc, df, table_name):
                    tables_processed += 1
                
                # 添加分页符(最后一个表除外)
                if tables.index(table_name) < len(tables) - 1:
                    doc.add_page_break()
            
            # 保存文档
            doc.save(output_file)
            print(f"成功导出 {tables_processed} 个表到 {output_file}")
    
    except Exception as e:
        print(f"错误: {e}")
        return False
    
    return True

def main():
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='将SQLite数据库中所有表导出到Word文档')
    parser.add_argument('--db', default=os.path.join(os.path.dirname(__file__), 'db.sqlite3'),
                        help='数据库文件路径')
    parser.add_argument('--output', default='数据库表导出.docx',
                        help='输出Word文件路径')
    parser.add_argument('--exclude', nargs='*', 
                        default=['sqlite_sequence', 'django_migrations'],
                        help='要排除的表名列表')
    args = parser.parse_args()

    try:
        # 导出所有表
        success = export_all_tables_to_docx(args.db, args.output, args.exclude)
        if success:
            print(f"导出成功：{args.output}")
            return 0
        else:
            return 1
    except Exception as e:
        print(f"错误: {e}")
        return 1

if __name__ == "__main__":
    main() 